from pathlib import Path
import re

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BASE = Path(__file__).resolve().parent
SOURCE = BASE / "岳声_简历_AI应用优化版.md"
OUTPUT = BASE / "岳声_简历_AI应用优化版_修正版.docx"


def set_font(style, name="Microsoft YaHei", size=None, bold=None, color=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = RGBColor(*color)


def shade(cell, fill="EAF2F8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bottom_border(paragraph, color="9EADBA", size="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_inline_bold(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part:
            paragraph.add_run(part)


def add_project_heading(doc, title, period):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(13.5)
    table.columns[1].width = Cm(4.0)
    left, right = table.rows[0].cells
    shade(left)
    shade(right)
    run = left.paragraphs[0].add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(31, 78, 121)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right.paragraphs[0].add_run(period)
    run.bold = True
    run.font.size = Pt(9)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)

    set_font(doc.styles["Normal"], size=9.2)
    doc.styles["Normal"].paragraph_format.space_after = Pt(1.5)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.06
    set_font(doc.styles["Title"], size=22, bold=True, color=(31, 78, 121))
    set_font(doc.styles["Heading 1"], size=12.5, bold=True, color=(31, 78, 121))
    set_font(doc.styles["Heading 2"], size=10.5, bold=True, color=(45, 45, 45))

    first_quote = True
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line[2:])
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2 if first_quote else 1)
            run = p.add_run(line[2:])
            if first_quote:
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(65, 65, 65)
                first_quote = False
            else:
                run.font.size = Pt(9)
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.add_run(line[3:])
            add_bottom_border(p)
        elif line.startswith("### "):
            content = line[4:]
            title, sep, period = content.rpartition(" | ")
            if sep:
                add_project_heading(doc, title, period)
            else:
                doc.add_paragraph(content, style="Heading 2")
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.48)
            p.paragraph_format.first_line_indent = Cm(-0.22)
            p.paragraph_format.space_after = Pt(1.2)
            add_inline_bold(p, line[2:])
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1.5)
            add_inline_bold(p, line)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("岳声 · Java 后端 / AI 应用开发")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(125, 125, 125)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
